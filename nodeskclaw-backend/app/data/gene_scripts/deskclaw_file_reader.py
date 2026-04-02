#!/usr/bin/env python3
"""DeskClaw File Reader Tool -- extract readable text from shared files."""

from __future__ import annotations

import argparse
import base64
import io
import subprocess
import sys
import tempfile
from pathlib import Path
from zipfile import BadZipFile

from _api_client import api_call, _output

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".xml", ".log"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(prog="deskclaw_file_reader", description="DeskClaw File Reader Tool")
    sub = parser.add_subparsers(dest="action", required=True)

    read = sub.add_parser("read", help="Read and extract text from a shared file")
    read.add_argument("--file-id", required=True)
    read.add_argument("--max-chars", type=int, default=12000)
    read.add_argument("--sheet-limit", type=int, default=5)
    read.add_argument("--row-limit", type=int, default=200)

    return parser


def _unwrap_data(resp: dict) -> dict:
    data = resp.get("data")
    if not isinstance(data, dict):
        raise ValueError("共享文件接口返回格式异常")
    return data


def _limit_text(text: str, max_chars: int) -> tuple[str, bool]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if len(normalized) <= max_chars:
        return normalized, False
    return normalized[:max_chars].rstrip(), True


def _read_text_bytes(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _read_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(f"[Page {index}]\n{page.extract_text() or ''}".strip())
    return "\n\n".join(pages)


def _read_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    lines: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_doc(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        result = subprocess.run(
            ["antiword", tmp.name],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        raise RuntimeError(stderr or "antiword 解析失败")
    return result.stdout


def _read_xlsx(content: bytes, sheet_limit: int, row_limit: int) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets[:sheet_limit]:
        parts.append(f"[Sheet] {sheet.title}")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                parts.append(" | ".join(values))
                row_count += 1
            if row_count >= row_limit:
                parts.append("... truncated rows ...")
                break
        parts.append("")
    return "\n".join(parts).strip()


def _read_xls(content: bytes, sheet_limit: int, row_limit: int) -> str:
    import xlrd

    workbook = xlrd.open_workbook(file_contents=content)
    parts: list[str] = []
    for sheet in workbook.sheets()[:sheet_limit]:
        parts.append(f"[Sheet] {sheet.name}")
        max_rows = min(sheet.nrows, row_limit)
        for row_idx in range(max_rows):
            values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
            if any(value.strip() for value in values):
                parts.append(" | ".join(values))
        if sheet.nrows > row_limit:
            parts.append("... truncated rows ...")
        parts.append("")
    return "\n".join(parts).strip()


def _extract_text(
    file_id: str,
    filename: str,
    content_type: str,
    content: bytes,
    *,
    sheet_limit: int,
    row_limit: int,
) -> dict:
    ext = Path(filename).suffix.lower()

    if ext in IMAGE_EXTENSIONS or content_type.startswith("image/"):
        url_resp = api_call("GET", f"/blackboard/files/{file_id}/url")
        url_data = _unwrap_data(url_resp)
        return {
            "mode": "model_vision",
            "filename": filename,
            "content_type": content_type,
            "message": "该文件属于图片，请直接把图片交给模型识别，不使用 OCR 脚本解析。",
            "url": url_data.get("url", ""),
        }

    if ext in TEXT_EXTENSIONS or content_type.startswith("text/"):
        return {"mode": "text", "text": _read_text_bytes(content)}
    if ext == ".pdf":
        return {"mode": "document", "text": _read_pdf(content)}
    if ext == ".docx":
        return {"mode": "document", "text": _read_docx(content)}
    if ext == ".doc":
        return {"mode": "document", "text": _read_doc(content)}
    if ext == ".xlsx":
        return {"mode": "spreadsheet", "text": _read_xlsx(content, sheet_limit, row_limit)}
    if ext == ".xls":
        return {"mode": "spreadsheet", "text": _read_xls(content, sheet_limit, row_limit)}

    raise ValueError(f"暂不支持读取该文件类型: {ext or content_type or 'unknown'}")


def main() -> None:
    parser = _build_parser()
    args = parser.parse_args()

    if args.action != "read":
        _output({"error": True, "message": f"不支持的 action: {args.action}"})
        sys.exit(1)

    resp = api_call("GET", f"/blackboard/files/{args.file_id}/content")
    payload = _unwrap_data(resp)
    filename = payload.get("filename") or args.file_id
    content_type = payload.get("content_type") or "application/octet-stream"
    content_b64 = payload.get("content") or ""
    content = base64.b64decode(content_b64)

    try:
        parsed = _extract_text(
            args.file_id,
            filename,
            content_type,
            content,
            sheet_limit=args.sheet_limit,
            row_limit=args.row_limit,
        )
    except ModuleNotFoundError as exc:
        _output({"error": True, "message": f"运行时缺少解析依赖: {exc.name}"})
        sys.exit(1)
    except BadZipFile:
        _output({"error": True, "message": f"文件损坏或格式不合法: {filename}"})
        sys.exit(1)
    except Exception as exc:
        _output({"error": True, "message": str(exc), "filename": filename, "content_type": content_type})
        sys.exit(1)

    if "text" in parsed:
        text, truncated = _limit_text(parsed["text"], args.max_chars)
        parsed["text"] = text
        parsed["truncated"] = truncated

    _output({
        "code": 0,
        "message": "success",
        "data": {
            "file_id": args.file_id,
            "filename": filename,
            "content_type": content_type,
            **parsed,
        },
    })


if __name__ == "__main__":
    main()
